"""
書類生成サービス

Word/Excelテンプレートへのデータ差し込みと書類生成
"""
import os
import re
import shutil
from datetime import datetime
from typing import Dict, Any, Optional, List, Tuple
from pathlib import Path

from utils.logger import get_logger

logger = get_logger(__name__)

# プレースホルダーのパターン: {{key}}
PLACEHOLDER_PATTERN = re.compile(r'\{\{(\w+)\}\}')


class DocumentServiceException(Exception):
    """書類サービス例外"""
    pass


class DocumentService:
    """
    書類生成サービス

    Word/Excelテンプレートにデータを差し込んで書類を生成
    """

    def __init__(self, output_dir: Optional[str] = None, template_dir: Optional[str] = None):
        """
        Args:
            output_dir: 出力先ディレクトリ（省略時はdata/documents/output）
            template_dir: テンプレート保存ディレクトリ（省略時はdata/documents/templates）
        """
        base_dir = Path(__file__).parent.parent / "data" / "documents"
        self.output_dir = Path(output_dir) if output_dir else base_dir / "output"
        self.template_dir = Path(template_dir) if template_dir else base_dir / "templates"

        # ディレクトリ作成
        self.output_dir.mkdir(parents=True, exist_ok=True)
        self.template_dir.mkdir(parents=True, exist_ok=True)

    def generate_document(
        self,
        template_path: str,
        data: Dict[str, Any],
        output_filename: Optional[str] = None
    ) -> str:
        """
        テンプレートからドキュメントを生成

        Args:
            template_path: テンプレートファイルのパス
            data: 差し込みデータ {key: value}
            output_filename: 出力ファイル名（省略時は自動生成）

        Returns:
            str: 生成されたファイルのパス

        Raises:
            DocumentServiceException: 生成に失敗した場合
        """
        template_path = Path(template_path)

        if not template_path.exists():
            raise DocumentServiceException(f"テンプレートファイルが見つかりません: {template_path}")

        ext = template_path.suffix.lower()

        if ext == ".docx":
            return self._generate_word_document(template_path, data, output_filename)
        elif ext == ".xlsx":
            return self._generate_excel_document(template_path, data, output_filename)
        else:
            raise DocumentServiceException(f"サポートされていないファイル形式です: {ext}")

    def _generate_word_document(
        self,
        template_path: Path,
        data: Dict[str, Any],
        output_filename: Optional[str] = None
    ) -> str:
        """Word文書を生成"""
        try:
            from docx import Document
            from docx.shared import Pt
        except ImportError:
            raise DocumentServiceException(
                "python-docxがインストールされていません。pip install python-docx を実行してください"
            )

        try:
            doc = Document(str(template_path))

            # 段落内のテキストを置換
            for paragraph in doc.paragraphs:
                self._replace_placeholder_in_paragraph(paragraph, data)

            # テーブル内のテキストを置換
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for paragraph in cell.paragraphs:
                            self._replace_placeholder_in_paragraph(paragraph, data)

            # ヘッダー・フッター内のテキストを置換
            for section in doc.sections:
                for header in [section.header, section.first_page_header, section.even_page_header]:
                    if header:
                        for paragraph in header.paragraphs:
                            self._replace_placeholder_in_paragraph(paragraph, data)
                for footer in [section.footer, section.first_page_footer, section.even_page_footer]:
                    if footer:
                        for paragraph in footer.paragraphs:
                            self._replace_placeholder_in_paragraph(paragraph, data)

            # 出力ファイル名
            if not output_filename:
                timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
                output_filename = f"{template_path.stem}_{timestamp}.docx"

            output_path = self.output_dir / output_filename
            doc.save(str(output_path))

            logger.info(f"Word document generated: {output_path}")
            return str(output_path)

        except Exception as e:
            logger.error(f"Failed to generate Word document: {e}")
            raise DocumentServiceException(f"Word文書の生成に失敗しました: {e}")

    def _replace_placeholder_in_paragraph(self, paragraph, data: Dict[str, Any]):
        """段落内のプレースホルダーを置換（書式を保持）"""
        # 段落全体のテキストを取得
        full_text = paragraph.text

        # プレースホルダーがなければスキップ
        if not PLACEHOLDER_PATTERN.search(full_text):
            return

        # プレースホルダーを置換
        new_text = PLACEHOLDER_PATTERN.sub(
            lambda m: str(data.get(m.group(1), m.group(0))),
            full_text
        )

        # 変更がなければスキップ
        if new_text == full_text:
            return

        # 書式を保持しながらテキストを更新
        # 最初のrunに全テキストを設定し、他のrunはクリア
        if paragraph.runs:
            first_run = paragraph.runs[0]
            first_run.text = new_text
            for run in paragraph.runs[1:]:
                run.text = ""
        else:
            paragraph.text = new_text

    def _generate_excel_document(
        self,
        template_path: Path,
        data: Dict[str, Any],
        output_filename: Optional[str] = None
    ) -> str:
        """Excelファイルを生成"""
        try:
            from openpyxl import load_workbook
        except ImportError:
            raise DocumentServiceException(
                "openpyxlがインストールされていません。pip install openpyxl を実行してください"
            )

        try:
            wb = load_workbook(str(template_path))

            for sheet in wb.worksheets:
                for row in sheet.iter_rows():
                    for cell in row:
                        if cell.value and isinstance(cell.value, str):
                            # プレースホルダーを置換
                            new_value = PLACEHOLDER_PATTERN.sub(
                                lambda m: str(data.get(m.group(1), m.group(0))),
                                cell.value
                            )
                            if new_value != cell.value:
                                cell.value = new_value

            # 出力ファイル名
            if not output_filename:
                timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
                output_filename = f"{template_path.stem}_{timestamp}.xlsx"

            output_path = self.output_dir / output_filename
            wb.save(str(output_path))

            logger.info(f"Excel document generated: {output_path}")
            return str(output_path)

        except Exception as e:
            logger.error(f"Failed to generate Excel document: {e}")
            raise DocumentServiceException(f"Excelファイルの生成に失敗しました: {e}")

    def save_template(self, source_path: str, name: str) -> Tuple[str, str]:
        """
        テンプレートファイルを保存

        Args:
            source_path: 元ファイルのパス
            name: テンプレート名

        Returns:
            Tuple[str, str]: (保存先パス, ファイル種別)
        """
        source = Path(source_path)

        if not source.exists():
            raise DocumentServiceException(f"ファイルが見つかりません: {source_path}")

        ext = source.suffix.lower()
        if ext not in [".docx", ".xlsx"]:
            raise DocumentServiceException(f"サポートされていないファイル形式です: {ext}")

        # 安全なファイル名に変換
        safe_name = re.sub(r'[^\w\-_]', '_', name)
        dest_filename = f"{safe_name}{ext}"
        dest_path = self.template_dir / dest_filename

        # コピー
        shutil.copy2(source, dest_path)

        file_type = ext[1:]  # .docx -> docx
        logger.info(f"Template saved: {dest_path}")

        return str(dest_path), file_type

    def extract_placeholders(self, template_path: str) -> List[str]:
        """
        テンプレートからプレースホルダーを抽出

        Args:
            template_path: テンプレートファイルのパス

        Returns:
            List[str]: プレースホルダーのキー一覧
        """
        template_path = Path(template_path)

        if not template_path.exists():
            raise DocumentServiceException(f"テンプレートファイルが見つかりません: {template_path}")

        ext = template_path.suffix.lower()
        placeholders = set()

        if ext == ".docx":
            placeholders = self._extract_placeholders_from_word(template_path)
        elif ext == ".xlsx":
            placeholders = self._extract_placeholders_from_excel(template_path)

        return sorted(list(placeholders))

    def _extract_placeholders_from_word(self, template_path: Path) -> set:
        """Wordファイルからプレースホルダーを抽出"""
        try:
            from docx import Document
        except ImportError:
            return set()

        placeholders = set()

        try:
            doc = Document(str(template_path))

            # 段落
            for paragraph in doc.paragraphs:
                matches = PLACEHOLDER_PATTERN.findall(paragraph.text)
                placeholders.update(matches)

            # テーブル
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for paragraph in cell.paragraphs:
                            matches = PLACEHOLDER_PATTERN.findall(paragraph.text)
                            placeholders.update(matches)

            # ヘッダー・フッター
            for section in doc.sections:
                for header in [section.header, section.first_page_header, section.even_page_header]:
                    if header:
                        for paragraph in header.paragraphs:
                            matches = PLACEHOLDER_PATTERN.findall(paragraph.text)
                            placeholders.update(matches)
                for footer in [section.footer, section.first_page_footer, section.even_page_footer]:
                    if footer:
                        for paragraph in footer.paragraphs:
                            matches = PLACEHOLDER_PATTERN.findall(paragraph.text)
                            placeholders.update(matches)

        except Exception as e:
            logger.error(f"Failed to extract placeholders from Word: {e}")

        return placeholders

    def _extract_placeholders_from_excel(self, template_path: Path) -> set:
        """Excelファイルからプレースホルダーを抽出"""
        try:
            from openpyxl import load_workbook
        except ImportError:
            return set()

        placeholders = set()

        try:
            wb = load_workbook(str(template_path), data_only=False)

            for sheet in wb.worksheets:
                for row in sheet.iter_rows():
                    for cell in row:
                        if cell.value and isinstance(cell.value, str):
                            matches = PLACEHOLDER_PATTERN.findall(cell.value)
                            placeholders.update(matches)

        except Exception as e:
            logger.error(f"Failed to extract placeholders from Excel: {e}")

        return placeholders

    def delete_template(self, template_path: str) -> bool:
        """
        テンプレートファイルを削除

        Args:
            template_path: テンプレートファイルのパス

        Returns:
            bool: 削除成功時True
        """
        try:
            path = Path(template_path)
            if path.exists():
                path.unlink()
                logger.info(f"Template deleted: {template_path}")
                return True
            return False
        except Exception as e:
            logger.error(f"Failed to delete template: {e}")
            return False

    def get_available_fields(self) -> Dict[str, List[Dict[str, str]]]:
        """
        差し込み可能なフィールド一覧を取得

        Returns:
            Dict: カテゴリ別のフィールド定義
        """
        return {
            "staff": [
                {"key": "staff_name", "label": "職員氏名", "source": "staff.name"},
                {"key": "staff_email", "label": "職員メール", "source": "staff.email"},
                {"key": "staff_phone", "label": "職員電話番号", "source": "staff.phone"},
                {"key": "staff_address", "label": "職員住所", "source": "staff.address"},
                {"key": "staff_type", "label": "職員種別", "source": "staff.staff_type"},
            ],
            "hospital": [
                {"key": "hospital_name", "label": "病院名", "source": "hospital.name"},
                {"key": "hospital_address", "label": "病院住所", "source": "hospital.address"},
                {"key": "hospital_phone", "label": "病院電話番号", "source": "hospital.phone"},
                {"key": "hospital_director", "label": "院長名", "source": "hospital.director_name"},
            ],
            "date": [
                {"key": "today", "label": "今日の日付", "source": "date.today"},
                {"key": "today_jp", "label": "今日の日付（和暦）", "source": "date.today_jp"},
                {"key": "fiscal_year", "label": "年度", "source": "date.fiscal_year"},
            ],
            "custom": [
                {"key": "custom_text", "label": "自由入力", "source": "custom"},
            ],
        }


def get_document_service(
    output_dir: Optional[str] = None,
    template_dir: Optional[str] = None
) -> DocumentService:
    """
    DocumentServiceのインスタンスを取得

    Args:
        output_dir: 出力先ディレクトリ
        template_dir: テンプレート保存ディレクトリ

    Returns:
        DocumentService: サービスインスタンス
    """
    return DocumentService(output_dir, template_dir)
